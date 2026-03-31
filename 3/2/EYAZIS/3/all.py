import streamlit as st
import pandas as pd
import json
import re
from io import BytesIO
import fitz
from docx import Document
from striprtf.striprtf import rtf_to_text

import nltk
from razdel import tokenize
import pymorphy3

# Импорт констант и таймера из внешних файлов
from constants import (
    MIN_YEAR, MAX_YEAR, DEFAULT_YEAR,
    POS_TRANSLATE, GRAMMEMES_TRANSLATE,
    TERMINOLOGY_TEXT, HELP_TEXT, MENU_ITEMS,
    SYNTAX_DEP_TRANSLATE, SYNTAX_POS_TRANSLATE
)
from performance_timer import PerformanceTimer

# Библиотеки для синтаксического анализа
import spacy
from spacy import displacy

# SQLAlchemy модели
from sqlalchemy import create_engine, Column, Integer, String, Text, func, select
from sqlalchemy.orm import sessionmaker, declarative_base
from sqlalchemy import text

nltk.download('punkt', quiet=True)

Base = declarative_base()


class CorpusDocument(Base):
    __tablename__ = 'documents'
    id = Column(Integer, primary_key=True)
    title = Column(String(255))
    author = Column(String(255))
    year = Column(Integer)
    type = Column(String(100))
    text = Column(Text)
    filename = Column(String(255))


class Sentence(Base):
    __tablename__ = 'sentences'
    id = Column(Integer, primary_key=True)
    doc_id = Column(Integer)
    sentence_index = Column(Integer)
    text = Column(Text)


class Token(Base):
    __tablename__ = 'tokens'
    id = Column(Integer, primary_key=True)
    sentence_id = Column(Integer)
    doc_id = Column(Integer)
    token_index = Column(Integer)
    wordform = Column(String(100))
    lemma = Column(String(100))
    pos_rus = Column(String(50))
    morph_rus = Column(Text)


class TagTranslator:
    POS_TRANSLATE = POS_TRANSLATE
    GRAMMEMES_TRANSLATE = GRAMMEMES_TRANSLATE

    def translate(self, tag_str):
        if not tag_str or tag_str == 'None':
            return ''
        parts = str(tag_str).replace(',', ', ').split()
        translated = []
        for part in parts:
            part = part.strip(', ')
            sub = [self.GRAMMEMES_TRANSLATE.get(p.strip(), "") for p in part.split(',')]
            translated.append(', '.join(sub))
        return ' '.join(translated).strip()

    def get_pos_rus(self, pos):
        return self.POS_TRANSLATE.get(pos, pos or "NONE")


class SyntaxTagTranslator:
    """Переводчик синтаксических тегов spaCy на русский язык"""
    
    DEP_TRANSLATE = SYNTAX_DEP_TRANSLATE
    POS_TRANSLATE = SYNTAX_POS_TRANSLATE

    def get_dep_rus(self, dep):
        """Перевод синтаксической зависимости на русский"""
        if not dep:
            return 'неизвестно'
        return self.DEP_TRANSLATE.get(dep, dep)

    def get_pos_rus(self, pos):
        """Перевод универсального POS тега на русский"""
        if not pos:
            return 'неизвестно'
        return self.POS_TRANSLATE.get(pos, pos)

    def translate_token(self, token_data):
        """Перевод всех синтаксических характеристик токена"""
        return {
            'text': token_data.get('text', ''),
            'lemma': token_data.get('lemma', ''),
            'pos_rus': self.get_pos_rus(token_data.get('pos', '')),
            'tag': token_data.get('tag', ''),
            'dep_rus': self.get_dep_rus(token_data.get('dep', '')),
            'dep': token_data.get('dep', ''),
            'head': token_data.get('head', ''),
            'head_idx': token_data.get('head_idx', 0)
        }


class FileHandler:
    def extract_text(self, uploaded_file):
        name = uploaded_file.name.lower()
        data = uploaded_file.getvalue()
        if name.endswith('.txt'):
            return data.decode('utf-8', errors='ignore')
        elif name.endswith('.pdf'):
            doc = fitz.open(stream=data, filetype="pdf")
            text = '\n'.join(page.get_text() or '' for page in doc)
            doc.close()
            return text
        elif name.endswith('.docx'):
            doc = Document(BytesIO(data))
            return '\n'.join(p.text for p in doc.paragraphs)
        elif name.endswith('.rtf') and rtf_to_text:
            return rtf_to_text(data.decode('utf-8', errors='ignore'))
        return None


class TextParser:
    def __init__(self):
        self.morph = pymorphy3.MorphAnalyzer(lang='ru')
        self.translator = TagTranslator()

    def analyze_text(self, text: str, doc_id: int):
        # Разбиваем текст на предложения (простая эвристика по концам предложений)
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        sentence_rows = []
        token_rows = []
        
        for sent_idx, sent_text in enumerate(sentences):
            # Добавляем предложение
            sentence_rows.append({
                'doc_id': doc_id,
                'sentence_index': sent_idx,
                'text': sent_text
            })
            
            # Токенизируем предложение
            tokens = [token.text for token in tokenize(sent_text) if re.match(r'^\w+$', token.text)]
            
            for token_idx, token in enumerate(tokens):
                if token.isdigit():
                    continue
                p = self.morph.parse(token)[0]
                pos = p.tag.POS or 'UNK'
                morph_tag = str(p.tag)
                token_rows.append({
                    'sentence_id': sent_idx,  # Временный ID, будет заменён на настоящий
                    'doc_id': doc_id,
                    'token_index': token_idx,
                    'wordform': token.lower(),
                    'lemma': p.normal_form.lower(),
                    'pos_rus': self.translator.get_pos_rus(pos),
                    'morph_rus': self.translator.translate(morph_tag),
                })
        
        return sentence_rows, token_rows


class DataStorage:
    def __init__(self):
        self.engine = create_engine(
            st.secrets['database']['url'],
            connect_args={'client_encoding': 'utf8'}
        )
        Base.metadata.create_all(self.engine)
        self.Session = sessionmaker(bind=self.engine)

    def add_sentences_and_tokens(self, sentence_rows, token_rows):
        """Добавление предложений и токенов в базу данных"""
        if sentence_rows:
            with self.Session() as session:
                # Сначала добавляем предложения
                session.bulk_insert_mappings(Sentence, sentence_rows)
                session.commit()
                
                # Получаем созданные ID предложений
                sentences = session.query(Sentence.id, Sentence.doc_id, Sentence.sentence_index).all()
                sentence_id_map = {(row.doc_id, row.sentence_index): row.id for row in sentences}
                
                # Обновляем sentence_id в токенах
                for token in token_rows:
                    temp_sent_id = token['sentence_id']
                    doc_id = token['doc_id']
                    real_sentence_id = sentence_id_map.get((doc_id, temp_sent_id))
                    if real_sentence_id:
                        token['sentence_id'] = real_sentence_id
                
                # Добавляем токены
                session.bulk_insert_mappings(Token, token_rows)
                session.commit()

    def add_document(self, text, metadata):
        doc = CorpusDocument(
            title=metadata['title'],
            author=metadata['author'],
            year=metadata['year'],
            type=metadata['type'],
            text=text,
            filename=metadata.get('filename')
        )
        with self.Session() as session:
            session.add(doc)
            session.commit()
            session.refresh(doc)
            return doc.id

    def update_document_text(self, doc_id, new_text):
        with self.Session() as session:
            doc = session.query(CorpusDocument).filter(CorpusDocument.id == doc_id).first()
            if doc:
                doc.text = new_text
                session.commit()

    def get_document_metadata(self, doc_id):
        with self.Session() as session:
            doc = session.query(CorpusDocument).filter(CorpusDocument.id == doc_id).first()
            if doc:
                return {
                    'title': doc.title,
                    'author': doc.author,
                    'year': doc.year,
                    'type': doc.type
                }
            return None

    def get_document_text(self, doc_id):
        with self.Session() as session:
            doc = session.query(CorpusDocument).filter(CorpusDocument.id == doc_id).first()
            return doc.text if doc else None

    def get_sentences_for_doc(self, doc_id):
        """Получение всех предложений для документа"""
        with self.Session() as session:
            sentences = session.query(Sentence).filter(Sentence.doc_id == doc_id)\
                .order_by(Sentence.sentence_index).all()
            return [{'id': s.id, 'text': s.text, 'sentence_index': s.sentence_index} for s in sentences]

    def get_tokens_for_sentence(self, sentence_id):
        """Получение всех токенов для предложения"""
        with self.Session() as session:
            tokens = session.query(Token).filter(Token.sentence_id == sentence_id)\
                .order_by(Token.token_index).all()
            return [{'wordform': t.wordform, 'lemma': t.lemma, 'pos_rus': t.pos_rus,
                     'morph_rus': t.morph_rus, 'token_index': t.token_index} for t in tokens]

    def remove_data_for_doc(self, doc_id):
        """Удаление токенов и предложений для документа"""
        with self.Session() as session:
            # Сначала удаляем токены (они ссылаются на предложения)
            session.query(Token).filter(Token.doc_id == doc_id).delete()
            session.commit()
            # Затем удаляем предложения
            session.query(Sentence).filter(Sentence.doc_id == doc_id).delete()
            session.commit()

    def get_doc_list(self):
        query = select(CorpusDocument.id, CorpusDocument.title)
        with self.engine.connect() as conn:
            result = conn.execute(query).fetchall()
        return {row[0]: {'metadata': {'title': row[1]}} for row in result}

    def get_top_wordforms(self, limit=20):
        query = "SELECT wordform, COUNT(*) as frequency FROM tokens GROUP BY wordform ORDER BY frequency DESC LIMIT :limit"
        with self.engine.connect() as conn:
            result = conn.execute(text(query), {'limit': limit}).fetchall()
        return pd.DataFrame(result, columns=['Словоформа', 'Частота'])

    def get_top_lemmas(self, limit=20):
        query = "SELECT lemma, COUNT(*) as frequency FROM tokens GROUP BY lemma ORDER BY frequency DESC LIMIT :limit"
        with self.engine.connect() as conn:
            result = conn.execute(text(query), {'limit': limit}).fetchall()
        return pd.DataFrame(result, columns=['Лемма', 'Частота'])

    def get_pos_distribution(self):
        query = "SELECT pos_rus, COUNT(*) as count FROM tokens GROUP BY pos_rus ORDER BY count DESC"
        with self.engine.connect() as conn:
            result = conn.execute(text(query)).fetchall()
        return pd.DataFrame(result, columns=['Часть речи', 'Количество'])

    def get_top_morph(self, limit=15):
        query = "SELECT morph_rus, COUNT(*) as frequency FROM tokens WHERE morph_rus != '' AND morph_rus IS NOT NULL GROUP BY morph_rus ORDER BY frequency DESC LIMIT :limit"
        with self.engine.connect() as conn:
            result = conn.execute(text(query), {'limit': limit}).fetchall()
        return pd.DataFrame(result, columns=['Морфологические характеристики', 'Частота'])

    def get_doc_stats(self):
        query = """
        SELECT d.title as doc_title, COUNT(t.id) as count
        FROM documents d
        LEFT JOIN tokens t ON d.id = t.doc_id
        GROUP BY d.title
        ORDER BY count DESC
        """
        with self.engine.connect() as conn:
            result = conn.execute(text(query)).fetchall()
        return pd.DataFrame(result, columns=['Название документа', 'Количество словоформ'])

    def search_words(self, query):
        PerformanceTimer.start()
        q = f"%{query.lower()}%"
        sql = """
        SELECT d.title as doc_title, t.wordform, t.lemma, t.pos_rus, t.morph_rus, d.author, d.year, d.type
        FROM tokens t
        JOIN documents d ON t.doc_id = d.id
        WHERE t.wordform ILIKE :q OR t.lemma ILIKE :q
        """
        with self.engine.connect() as conn:
            result = conn.execute(text(sql), {'q': q}).fetchall()
        PerformanceTimer.stop("Поиск слова:")
        return pd.DataFrame(result, columns=['Название документа', 'Словоформа', 'Лемма', 'Часть речи',
                                             'Морфологические характеристики', 'Автор', 'Год', 'Тип текста'])

    def search_phrases(self, phrase):
        """Поиск фразы в предложениях корпуса."""
        sql = """
        SELECT s.id, s.doc_id, s.text, d.title
        FROM sentences s
        JOIN documents d ON s.doc_id = d.id
        WHERE s.text ILIKE :q
        ORDER BY s.doc_id, s.sentence_index
        """
        with self.engine.connect() as conn:
            results = conn.execute(text(sql), {'q': f"%{phrase}%"}).fetchall()
        return results

    def get_filtered_data(self, author=None, doc_type=None, year=None):
        where_clauses = []
        params = {}
        if author and author != "Все":
            where_clauses.append("d.author = :author")
            params['author'] = author
        if doc_type and doc_type != "Все":
            where_clauses.append("d.type = :type")
            params['type'] = doc_type
        if year and year != "Все":
            where_clauses.append("d.year = :year")
            params['year'] = year

        where = "WHERE " + " AND ".join(where_clauses) if where_clauses else ""
        sql = f"""
        SELECT d.title as doc_title, t.wordform, t.lemma, t.pos_rus, t.morph_rus, d.author, d.year, d.type
        FROM tokens t
        JOIN documents d ON t.doc_id = d.id
        {where}
        """
        with self.engine.connect() as conn:
            result = conn.execute(text(sql), params).fetchall()
        return pd.DataFrame(result, columns=['Название документа', 'Словоформа', 'Лемма', 'Часть речи',
                                             'Морфологические характеристики', 'Автор', 'Год', 'Тип текста'])

    def get_authors(self):
        query = select(func.distinct(CorpusDocument.author))
        with self.engine.connect() as conn:
            result = conn.execute(query).fetchall()
        return sorted([row[0] for row in result if row[0]])

    def get_types(self):
        query = select(func.distinct(CorpusDocument.type))
        with self.engine.connect() as conn:
            result = conn.execute(query).fetchall()
        return sorted([row[0] for row in result if row[0]])

    def get_years(self):
        query = select(func.distinct(CorpusDocument.year))
        with self.engine.connect() as conn:
            result = conn.execute(query).fetchall()
        return sorted([row[0] for row in result if row[0]], reverse=True)

    def save_to_json(self):
        with self.Session() as session:
            docs = session.query(CorpusDocument).all()
            sentences = session.query(Sentence).all()
            tokens = session.query(Token).all()

        return json.dumps({
            'documents': [{'title': d.title, 'author': d.author, 'year': d.year,
                           'type': d.type, 'text': d.text, 'filename': d.filename}
                          for d in docs],
            'sentences': [{'doc_id': s.doc_id, 'sentence_index': s.sentence_index,
                           'text': s.text} for s in sentences],
            'tokens': [{'sentence_id': t.sentence_id, 'doc_id': t.doc_id,
                        'token_index': t.token_index, 'wordform': t.wordform,
                        'lemma': t.lemma, 'pos_rus': t.pos_rus, 'morph_rus': t.morph_rus}
                       for t in tokens]
        }, ensure_ascii=False, indent=2).encode('utf-8')

    def load_from_json(self, data):
        loaded = json.loads(data)

        with self.Session() as session:
            session.query(Token).delete()
            session.query(Sentence).delete()
            session.query(CorpusDocument).delete()
            session.commit()

            doc_id_map = {}
            for doc_dict in loaded['documents']:
                doc = CorpusDocument(**doc_dict)
                session.add(doc)
                session.flush()
                doc_id_map[doc.id] = doc.id

            session.commit()

            for sent_dict in loaded.get('sentences', []):
                sentence = Sentence(**sent_dict)
                session.add(sentence)
                session.flush()

            session.commit()

            for token_dict in loaded.get('tokens', []):
                token = Token(**token_dict)
                session.add(token)

            session.commit()


class SyntaxAnalyzer:
    def __init__(self):
        self.nlp = None
        self.model_name = "ru_core_news_sm"
        self.translator = SyntaxTagTranslator()

    def ensure_loaded(self):
        if self.nlp is None:
            try:
                self.nlp = spacy.load(self.model_name)
            except OSError:
                st.error(f"Модель {self.model_name} не найдена. Установите: python -m spacy download {self.model_name}")
                return False
        return True

    def analyze(self, text):
        if not self.ensure_loaded():
            return None
        return self.nlp(text)

    def get_sentence_trees(self, doc):
        sentences = []
        for sent in doc.sents:
            sent_doc = doc[sent.start:sent.end]
            # Визуализация дерева зависимостей
            html = displacy.render(sent_doc, style="dep", options={"compact": True, "bg": "#fff"}, jupyter=False, page=False)
            tokens_data = []
            for token in sent_doc:
                token_data = {
                    'text': token.text,
                    'lemma': token.lemma_,
                    'pos': token.pos_,
                    'tag': token.tag_,
                    'dep': token.dep_,
                    'head': token.head.text,
                    'head_idx': token.head.i - sent.start,
                }
                # Перевод синтаксических характеристик на русский
                translated = self.translator.translate_token(token_data)
                tokens_data.append(translated)
            sentences.append({
                'text': sent.text,
                'tokens': tokens_data,
                'html': html
            })
        return sentences


class CorpusManager:
    def __init__(self):
        self.storage = DataStorage()
        self.file_handler = FileHandler()
        self.parser = TextParser()
        self.view = View(self)
        self.syntax_analyzer = SyntaxAnalyzer()

    def add_document(self, text: str, metadata: dict):
        doc_id = self.storage.add_document(text, metadata)
        sentence_rows, token_rows = self.parser.analyze_text(text, doc_id)
        self.storage.add_sentences_and_tokens(sentence_rows, token_rows)
        return doc_id

    def update_document(self, doc_id: int, new_text: str):
        metadata = self.storage.get_document_metadata(doc_id)
        if metadata:
            self.storage.remove_data_for_doc(doc_id)
            sentence_rows, token_rows = self.parser.analyze_text(new_text, doc_id)
            self.storage.add_sentences_and_tokens(sentence_rows, token_rows)
            self.storage.update_document_text(doc_id, new_text)

    def process_uploaded_files(self, uploaded_files, author, title_base, year, doc_type):
        PerformanceTimer.start()
        for file in uploaded_files:
            text = self.file_handler.extract_text(file)
            if text:
                metadata = {
                    "title": f"{title_base} — {file.name}",
                    "author": author,
                    "year": year,
                    "type": doc_type,
                    "filename": file.name
                }
                doc_id = self.add_document(text, metadata)
                st.success(f"{file.name} добавлен (ID {doc_id})")
        PerformanceTimer.stop()

    def handle_menu_selection(self, menu):
        if menu == "Загрузка и построение корпуса":
            self.handle_upload()
        elif menu == "Просмотр и редактирование корпуса":
            self.handle_view_edit()
        elif menu == "Поиск и анализ":
            self.handle_search_analysis()
        elif menu == "Синтаксический анализ":
            self.handle_syntax_analysis()
        elif menu == "Сохранение / Загрузка":
            self.handle_save_load()
        elif menu == "Терминологическая справка":
            self.handle_terminology()
        else:
            self.handle_help()
        self.view.render_sidebar_caption()

    def handle_upload(self):
        uploaded_files, author, title_base, year, doc_type, process_button = self.view.render_upload()
        if uploaded_files and process_button:
            self.process_uploaded_files(uploaded_files, author, title_base, year, doc_type)

    def handle_view_edit(self):
        selected_id, edited_text, save_button = self.view.render_view_edit()
        if save_button and selected_id is not None:
            self.update_document(selected_id, edited_text)
            st.success("Изменения сохранены и корпус обновлён")

    def handle_search_analysis(self):
        self.view.render_search_analysis()

    def handle_syntax_analysis(self):
        self.view.render_syntax_analysis()

    def handle_save_load(self):
        save_button, uploaded_json = self.view.render_save_load()
        if save_button:
            data = self.save_corpus_to_json()
            self.view.display_saved_json(data)
        if uploaded_json:
            data = uploaded_json.getvalue().decode('utf-8')
            self.load_corpus_from_json(data)
            st.success("Корпус успешно загружен!")

    def handle_terminology(self):
        self.view.render_terminology()

    def handle_help(self):
        self.view.render_help()

    def run(self):
        self.view.display_menu()

    def get_doc_text(self, doc_id: str) -> str:
        return self.storage.get_document_text(doc_id)

    def has_documents(self) -> bool:
        try:
            with self.storage.engine.connect() as conn:
                count = conn.execute(text("SELECT COUNT(*) FROM documents")).scalar()
                return count > 0
        except Exception as e:
            print(f"Ошибка при проверке наличия документов: {e}")
            return False

    def get_phrase_context(self, text: str, phrase: str, words_count: int = 80):
        if not phrase or not text:
            return []

        results = []
        text_lower = text.lower()
        phrase_lower = phrase.lower()
        start_pos = 0

        while True:
            pos = text_lower.find(phrase_lower, start_pos)
            if pos == -1:
                break

            before_text = text[:pos]
            after_text = text[pos + len(phrase):]

            words_before = re.findall(r'\w+', before_text, re.UNICODE)
            words_before = words_before[-words_count:] if len(words_before) > words_count else words_before

            words_after = re.findall(r'\w+', after_text, re.UNICODE)
            words_after = words_after[:words_count] if len(words_after) > words_count else words_after

            if words_before:
                first_word = words_before[0]
                left_start = before_text.rfind(first_word)
                left_context = text[left_start:pos]
            else:
                left_context = ""

            if words_after:
                last_word = words_after[-1]
                right_end = after_text.find(last_word) + len(last_word)
                right_context = after_text[:right_end]
            else:
                right_context = ""

            context = {
                'left': left_context.strip(),
                'phrase': text[pos:pos + len(phrase)],
                'right': right_context.strip()
            }
            results.append(context)

            start_pos = pos + 1

        return results

    def get_doc_list(self):
        return self.storage.get_doc_list()

    def get_top_wordforms(self, limit=20):
        return self.storage.get_top_wordforms(limit)

    def get_top_lemmas(self, limit=20):
        return self.storage.get_top_lemmas(limit)

    def get_pos_distribution(self):
        return self.storage.get_pos_distribution()

    def get_top_morph(self, limit=15):
        return self.storage.get_top_morph(limit)

    def get_doc_stats(self):
        return self.storage.get_doc_stats()

    def search_words(self, query):
        return self.storage.search_words(query)

    def search_phrases(self, phrase):
        return self.storage.search_phrases(phrase)

    def get_filtered_data(self, author=None, doc_type=None, year=None):
        return self.storage.get_filtered_data(author, doc_type, year)

    def get_authors(self):
        return self.storage.get_authors()

    def get_types(self):
        return self.storage.get_types()

    def get_years(self):
        return self.storage.get_years()

    def save_corpus_to_json(self):
        return self.storage.save_to_json()

    def load_corpus_from_json(self, data):
        self.storage.load_from_json(data)

    # Методы для работы с предложениями
    def get_sentences_for_doc(self, doc_id):
        return self.storage.get_sentences_for_doc(doc_id)

    def get_tokens_for_sentence(self, sentence_id):
        return self.storage.get_tokens_for_sentence(sentence_id)

    # Методы синтаксического анализа
    def analyze_document_syntax(self, doc_id):
        text = self.storage.get_document_text(doc_id)
        if not text:
            return None
        doc = self.syntax_analyzer.analyze(text)
        if doc is None:
            return None
        return self.syntax_analyzer.get_sentence_trees(doc)

    def analyze_text_syntax(self, text):
        doc = self.syntax_analyzer.analyze(text)
        if doc is None:
            return None
        return self.syntax_analyzer.get_sentence_trees(doc)


class View:
    def __init__(self, corpus_manager):
        self.corpus_manager = corpus_manager
        st.set_page_config(page_title="Корпусный менеджер", layout="wide")
        st.title("Корпусный менеджер текстов")

    def display_menu(self):
        # Инициализация состояния меню
        if "menu_selection" not in st.session_state:
            st.session_state.menu_selection = MENU_ITEMS[0]
        
        menu = st.sidebar.selectbox(
            "Меню",
            MENU_ITEMS,
            key="menu_selection"
        )
        self.corpus_manager.handle_menu_selection(menu)

    def render_upload(self):
        st.header("Загрузка текстов в корпус")
        uploaded_files = st.file_uploader(
            "Выберите файлы (TXT, PDF, DOCX, RTF)",
            accept_multiple_files=True,
            type=['txt', 'pdf', 'docx', 'rtf']
        )

        st.subheader("Метаданные")
        col1, col2 = st.columns(2)
        with col1:
            author = st.text_input("Автор", "Неизвестен", key="author_input")
            title_base = st.text_input("Базовое название", "Документ", key="title_input")
        with col2:
            year = st.number_input("Год", MIN_YEAR, MAX_YEAR, DEFAULT_YEAR, key="year_input")
            doc_type = st.selectbox("Тип текста", ["Художественный", "Научный", "Публицистический", "Другой"], key="doctype_input")

        process_button = st.button(" Обработать и добавить в корпус", type="primary", key="process_files_button")
        return uploaded_files, author, title_base, year, doc_type, process_button

    def render_view_edit(self):
        st.header("Просмотр и редактирование корпуса")
        docs = self.corpus_manager.get_doc_list()

        if not docs:
            st.info("Корпус пока пустой")
            return None, None, False

        selected_id = st.selectbox(
            "Выберите документ",
            list(docs.keys()),
            format_func=lambda x: f"{x} — {docs[x]['metadata']['title']}"
        )

        text = self.corpus_manager.get_doc_text(selected_id)

        edited_text = st.text_area("Редактировать текст", text, height=400)
        save_button = st.button("Сохранить изменения", key="save_changes_button")

        return selected_id, edited_text, save_button

    def render_search_analysis(self):
        st.header("Поиск и анализ корпуса")

        if not self.corpus_manager.has_documents():
            st.warning("Сначала загрузите документы в раздел «Загрузка»")
            return

        # Инициализация состояния вкладок
        if "search_tab" not in st.session_state:
            st.session_state.search_tab = 0
        
        tab_names = [
            "Общая статистика корпуса",
            "Поиск словоформ и лемм",
            "Поиск по фразе / кускам текста",
            "Фильтры по метаданным"
        ]
        
        tabs = st.tabs(tab_names)
        
        with tabs[0]:
            self.display_stats()

        with tabs[1]:
            query = self.render_word_search()
            self.display_word_search_results(query)

        with tabs[2]:
            phrase = self.render_phrase_search()
            self.display_phrase_search_results(phrase)

        with tabs[3]:
            self.render_filters()

    def display_stats(self):
        st.subheader("Частотные характеристики по всему корпусу")

        col1, col2 = st.columns(2)
        with col1:
            st.write("**Топ-20 словоформ**")
            wf = self.corpus_manager.get_top_wordforms()
            st.dataframe(wf, use_container_width=True)

            st.write("**Топ-20 лемм**")
            lem = self.corpus_manager.get_top_lemmas()
            st.dataframe(lem, use_container_width=True)

        with col2:
            st.write("**Распределение по частям речи**")
            pos_df = self.corpus_manager.get_pos_distribution()
            st.dataframe(pos_df, use_container_width=True)

        st.write("**Морфологические характеристики (топ-15)**")
        morph_df = self.corpus_manager.get_top_morph()
        st.dataframe(morph_df, use_container_width=True)

        st.write("**Статистика по документам**")
        doc_stat = self.corpus_manager.get_doc_stats()
        st.dataframe(doc_stat, use_container_width=True)

    def render_word_search(self):
        query = st.text_input("Введите слово или лемму для поиска")
        return query

    def display_word_search_results(self, query):
        if query:
            result = self.corpus_manager.search_words(query)
            st.dataframe(result, use_container_width=True)

    def render_phrase_search(self):
        st.subheader("Поиск по произвольному тексту / словосочетанию")
        phrase = st.text_input("Введите фразу или кусок текста", placeholder="Пользовательский интерфейс")
        return phrase

    def display_phrase_search_results(self, phrase: str):
        PerformanceTimer.start()
        if not phrase:
            return

        st.write(f"**Результаты поиска:** `{phrase}`")

        # Поиск в предложениях корпуса
        results = self.corpus_manager.search_phrases(phrase)

        if not results:
            st.info("Фраза не найдена ни в одном документе корпуса.")
            return

        match_counter = 0
        current_doc = None

        for sent_id, doc_id, sent_text, doc_title in results:
            doc_title_display = doc_title or f"Документ {doc_id}"
            
            # Показываем заголовок документа только при смене документа
            if current_doc != doc_id:
                st.subheader(f"{doc_title_display}")
                current_doc = doc_id
            
            match_counter += 1
            with st.expander(f"Находка #{match_counter}: {sent_text[:80]}..."):
                st.markdown(f"**{sent_text}**")

        if match_counter == 0:
            st.info("Фраза не найдена ни в одном документе корпуса.")
        PerformanceTimer.stop("Поиск фразы:")

    def render_filters(self):
        st.subheader("Фильтрация по метаданным")
        authors = ["Все"] + self.corpus_manager.get_authors()
        types_list = ["Все"] + self.corpus_manager.get_types()
        years = ["Все"] + self.corpus_manager.get_years()

        # Инициализация состояния фильтров
        if "filter_author" not in st.session_state:
            st.session_state.filter_author = "Все"
        if "filter_type" not in st.session_state:
            st.session_state.filter_type = "Все"
        if "filter_year" not in st.session_state:
            st.session_state.filter_year = "Все"

        col1, col2, col3 = st.columns(3)
        with col1:
            sel_author = st.selectbox("Автор", authors, key="filter_author")
        with col2:
            sel_type = st.selectbox("Тип текста", types_list, key="filter_type")
        with col3:
            sel_year = st.selectbox("Год", years, key="filter_year")

        apply_button = st.button("Применить фильтры", type="primary", key="apply_filters_button")

        if apply_button:
            filtered = self.corpus_manager.get_filtered_data(
                st.session_state.filter_author,
                st.session_state.filter_type,
                st.session_state.filter_year
            )
            st.dataframe(filtered, use_container_width=True)

    def render_syntax_analysis(self):
        st.header("Синтаксический анализ текста")

        if not self.corpus_manager.has_documents():
            st.warning("Сначала загрузите документы в корпус.")
            st.subheader("Или введите текст для анализа:")
            input_text = st.text_area("Введите текст для синтаксического разбора", height=200)
            if st.button("Анализировать введённый текст"):
                if input_text.strip():
                    with st.spinner("Выполняется синтаксический анализ..."):
                        sentences = self.corpus_manager.analyze_text_syntax(input_text)
                        self.display_syntax_results(sentences)
            return

        # Выбор документа из корпуса
        docs = self.corpus_manager.get_doc_list()
        selected_id = st.selectbox(
            "Выберите документ для анализа",
            list(docs.keys()),
            format_func=lambda x: f"{x} — {docs[x]['metadata']['title']}"
        )

        analyze_button = st.button("Выполнить синтаксический анализ документа")

        if analyze_button and selected_id is not None:
            with st.spinner("Выполняется синтаксический анализ..."):
                sentences = self.corpus_manager.analyze_document_syntax(selected_id)
                self.display_syntax_results(sentences)

    def display_syntax_results(self, sentences):
        if sentences is None:
            st.error("Ошибка при анализе. Убедитесь, что установлена модель spaCy.")
            return
        if not sentences:
            st.info("Нет данных для отображения.")
            return

        st.success(f"Анализ завершён. Обработано предложений: {len(sentences)}")

        # Сохранение результатов в JSON
        if st.button("Экспортировать результаты в JSON"):
            result_data = []
            for sent in sentences:
                result_data.append({
                    'text': sent['text'],
                    'tokens': sent['tokens']
                })
            json_str = json.dumps(result_data, ensure_ascii=False, indent=2)
            st.download_button(
                label="Скачать JSON",
                data=json_str.encode('utf-8'),
                file_name="syntax_analysis.json",
                mime="application/json"
            )

        # Отображение по предложениям с русскими названиями колонок
        for i, sent in enumerate(sentences):
            with st.expander(f"Предложение {i+1}: {sent['text'][:100]}..."):
                st.components.v1.html(sent['html'], height=300, scrolling=True)
                # Создаём DataFrame с русскими названиями колонок
                df_tokens = pd.DataFrame(sent['tokens'])
                df_display = pd.DataFrame({
                    'Слово': df_tokens['text'],
                    'Лемма': df_tokens['lemma'],
                    'Часть речи': df_tokens['pos_rus'],
                    'Тег': df_tokens['tag'],
                    'Синтаксическая роль': df_tokens['dep_rus'],
                    'Зависимость': df_tokens['dep'],
                    'Главное слово': df_tokens['head']
                })
                st.dataframe(df_display, use_container_width=True)

    def render_save_load(self):
        st.header("Сохранение и загрузка корпуса")
        col1, col2 = st.columns(2)
        save_button = False
        uploaded_json = None
        with col1:
            save_button = st.button("Сохранить корпус как JSON", key="save_json_button")
        with col2:
            uploaded_json = st.file_uploader("Загрузить сохранённый корпус", type='json')
        return save_button, uploaded_json

    def display_saved_json(self, data):
        st.download_button(
            label="Скачать corpus.json",
            data=data,
            file_name="corpus.json",
            mime="application/json",
            key="download_corpus"
        )

    def render_help(self):
        st.header("Система помощи")
        st.markdown(HELP_TEXT)

    def render_terminology(self):
        st.header("Терминологическая справка")
        st.markdown(TERMINOLOGY_TEXT)

    def render_sidebar_caption(self):
        st.sidebar.caption("Корпусный менеджер v0.999")


if __name__ == "__main__":
    manager = CorpusManager()
    manager.run()